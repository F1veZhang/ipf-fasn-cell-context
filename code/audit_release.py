from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from lxml import etree


WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": WORD_NS}


def citation_numbers(text: str) -> set[int]:
    numbers: set[int] = set()
    for match in re.finditer(r"\[(\d+(?:[-,]\d+)*)\]", text):
        for part in match.group(1).split(","):
            if "-" in part:
                start, end = map(int, part.split("-"))
                numbers.update(range(start, end + 1))
            else:
                numbers.add(int(part))
    return numbers


def sequence_fields(path: Path) -> tuple[dict[str, list[int]], list[str]]:
    sequences: dict[str, list[int]] = {"Figure": [], "Table": []}
    captions: list[str] = []
    with ZipFile(path) as archive:
        root = etree.fromstring(archive.read("word/document.xml"))

    for paragraph in root.xpath("//w:p[.//w:fldSimple]", namespaces=NS):
        found_sequence = False
        for field in paragraph.xpath(".//w:fldSimple", namespaces=NS):
            instruction = field.get(f"{{{WORD_NS}}}instr", "")
            match = re.search(r"\bSEQ\s+(Figure|Table)\b", instruction)
            if not match:
                continue
            cached_text = "".join(field.xpath(".//w:t/text()", namespaces=NS)).strip()
            if not cached_text.isdigit():
                raise ValueError(f"Non-numeric cached SEQ result: {cached_text!r}")
            sequences[match.group(1)].append(int(cached_text))
            found_sequence = True
        if found_sequence:
            captions.append("".join(paragraph.xpath(".//w:t/text()", namespaces=NS)))
    return sequences, captions


def main() -> None:
    parser = argparse.ArgumentParser(description="Audit the v1.3.1 manuscript release.")
    parser.add_argument("manuscript", type=Path)
    args = parser.parse_args()

    document = Document(args.manuscript)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    references = [
        paragraph for paragraph in document.paragraphs if paragraph.style.name == "Reference"
    ]
    cited = citation_numbers(text)
    sequences, materialized_captions = sequence_fields(args.manuscript)
    report = {
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "inline_shapes": len(document.inline_shapes),
        "references": len(references),
        "cited_unique": sorted(cited),
        "uncited_references": sorted(set(range(1, len(references) + 1)) - cited),
        "prespecified_hits": text.lower().count("prespecified"),
        "v1_2_hits": text.count("v1.2"),
        "old_figure_1_order_hits": text.count(
            "before transcriptome-wide whole-lung calibration"
        ),
        "quasibinomial_hits": text.lower().count("quasibinomial"),
        "sequence_fields": sequences,
        "materialized_captions": materialized_captions,
    }
    print(json.dumps(report, indent=2))

    if report["uncited_references"]:
        raise SystemExit("Uncited references remain")
    if report["prespecified_hits"] or report["v1_2_hits"] or report["old_figure_1_order_hits"]:
        raise SystemExit("Stale manuscript wording remains")
    if len(references) != 29 or len(document.inline_shapes) != 4 or len(document.tables) != 1:
        raise SystemExit("Unexpected manuscript structure")
    if sequences != {"Figure": [1, 2, 3, 4], "Table": [1]}:
        raise SystemExit(f"Unexpected Figure/Table sequence fields: {sequences}")


if __name__ == "__main__":
    main()
