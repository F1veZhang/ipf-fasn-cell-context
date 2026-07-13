from __future__ import annotations

import html
import json
import math
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE_DOC = None
REFERENCE_JSON = None
OUTPUT = None
ACTION_OUTPUT = None
FIGURES = {}

BODY_FONT = "Times New Roman"
ACCENT = RGBColor(35, 78, 112)
MUTED = RGBColor(85, 96, 106)
LIGHT_FILL = "EAF1F6"
HEADER_FILL = "DCE8F2"


def set_font(run, name=BODY_FONT, size=11, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=80, bottom=80, end=80):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_in):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths_dxa = [int(width * 1440) for width in widths_in]
    total = sum(widths_dxa)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_in[index])
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text_run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "1"
    text_run.append(text)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text_run, end])


def add_seq_field(paragraph, label, number, bookmark_name):
    bookmark_id = str(100 + number + (0 if label == "Figure" else 50))
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bookmark_id)
    start.set(qn("w:name"), bookmark_name)
    paragraph._p.append(start)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), f"SEQ {label} \\* ARABIC")
    run = OxmlElement("w:r")
    run_pr = OxmlElement("w:rPr")
    bold = OxmlElement("w:b")
    run_pr.append(bold)
    run.append(run_pr)
    text = OxmlElement("w:t")
    text.text = str(number)
    run.append(text)
    field.append(run)
    paragraph._p.append(field)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bookmark_id)
    paragraph._p.append(end)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.space_after = Pt(5)

    for style_name, size, color, before, after in (
        ("Heading 1", 14, ACCENT, 12, 5),
        ("Heading 2", 11.5, RGBColor(31, 48, 61), 9, 3),
        ("Heading 3", 10.8, RGBColor(31, 48, 61), 7, 2),
    ):
        style = doc.styles[style_name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    caption.font.name = BODY_FONT
    caption._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    caption.font.size = Pt(9)
    caption.font.bold = False
    caption.font.color.rgb = RGBColor(45, 52, 58)
    caption.paragraph_format.line_spacing = 1.0
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(8)

    if "Reference" not in [style.name for style in doc.styles]:
        reference = doc.styles.add_style("Reference", WD_STYLE_TYPE.PARAGRAPH)
    else:
        reference = doc.styles["Reference"]
    reference.font.name = BODY_FONT
    reference._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    reference._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    reference.font.size = Pt(9)
    reference.paragraph_format.left_indent = Cm(0.65)
    reference.paragraph_format.first_line_indent = Cm(-0.65)
    reference.paragraph_format.space_after = Pt(2)
    reference.paragraph_format.line_spacing = 1.0

    for sec in doc.sections:
        header = sec.header.paragraphs[0]
        header.text = "Cell-contextual metabolic remodeling in IPF"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in header.runs:
            set_font(run, size=8.5, italic=True, color=MUTED)
        footer = sec.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_field(footer)
        for run in footer.runs:
            set_font(run, size=8.5, color=MUTED)

    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def add_body(doc, text, bold_lead=None):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.widow_control = True
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_font(lead, bold=True)
        rest = paragraph.add_run(text[len(bold_lead):])
        set_font(rest)
    else:
        run = paragraph.add_run(text)
        set_font(run)
    return paragraph


def add_labeled_paragraph(doc, label, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    lead = paragraph.add_run(label + " ")
    set_font(lead, bold=True)
    body = paragraph.add_run(text)
    set_font(body)
    return paragraph


def add_figure(doc, number, caption_text, width=6.55):
    path = FIGURES[number]
    if not path.exists():
        raise FileNotFoundError(path)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    shape._inline.docPr.set("descr", f"Figure {number}: {caption_text[:180]}")
    shape._inline.docPr.set("title", f"Figure {number}")
    caption = doc.add_paragraph(style="Caption")
    caption.paragraph_format.keep_together = True
    lead = caption.add_run("Figure ")
    set_font(lead, size=9, bold=True)
    add_seq_field(caption, "Figure", number, f"fig{number}")
    dot = caption.add_run(". ")
    set_font(dot, size=9, bold=True)
    body = caption.add_run(caption_text)
    set_font(body, size=9, bold=False)


def add_table_caption(doc, number, text):
    caption = doc.add_paragraph(style="Caption")
    caption.paragraph_format.keep_with_next = True
    lead = caption.add_run("Table ")
    set_font(lead, size=9, bold=True)
    add_seq_field(caption, "Table", number, f"tbl{number}")
    dot = caption.add_run(". ")
    set_font(dot, size=9, bold=True)
    body = caption.add_run(text)
    set_font(body, size=9, bold=False)


def add_dataset_table(doc):
    add_table_caption(doc, 1, "Data sources, sample composition, and analytical roles.")
    headers = ["Source", "Modality", "Sample composition", "Primary role", "Interpretive boundary"]
    rows = [
        ["FinnGen R12 IPF", "GWAS summary statistics", "2,700 cases; 494,583 controls", "IPF outcome for nominal genetic screen", "Exposure provenance and full testing universe were unavailable in the project archive"],
        ["GSE245965", "Purified AT2 RNA-seq", "2 IPF; 3 normal", "AT2 differential expression and GSEA", "Small epithelial-context dataset; exploratory"],
        ["GSE150910", "Whole-lung RNA-seq", "103 IPF; 103 controls", "Standardized replication and composition sensitivity", "Bulk tissue; cell composition can influence effects"],
        ["GSE110147", "Whole-lung microarray", "22 IPF; 11 controls", "Standardized replication and composition sensitivity", "Small cohort; adjusted models have wide confidence intervals"],
        ["GSE24206", "Whole-lung microarray", "17 IPF; 6 healthy", "Standardized replication; exploratory stage context", "Small cohort and mixed disease stages"],
        ["GSE136831", "Lung scRNA-seq", "28 control and 32 IPF donors; 230,550 cells after study filters", "Cell-type localization; donor summaries; NNLS reference", "Dissociation-based; no spatial information"],
        ["GSE233844", "PBMC scRNA-seq", "13 controls; 13 stable IPF; 12 progressive IPF; 149,564 cells", "Supplementary systemic disease layer", "Not used as lung-local or prognostic evidence"],
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_repeat_table_header(table.rows[0])
    for index, value in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = value
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, HEADER_FILL)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                set_font(run, size=8.2, bold=True, color=RGBColor(25, 45, 60))
    for row_index, values in enumerate(rows, start=1):
        cells = table.add_row().cells
        for column_index, value in enumerate(values):
            cell = cells[column_index]
            cell.text = value
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if row_index % 2 == 0:
                set_cell_shading(cell, "F7F9FB")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    set_font(run, size=7.8)
    set_table_geometry(table, [0.95, 0.95, 1.25, 1.45, 1.85])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def clean_markup(value):
    value = html.unescape(value or "")
    value = re.sub(r"<[^>]+>", " ", value)
    value = re.sub(r"\s+", " ", value).strip()
    return value


def initials(given):
    parts = re.findall(r"[A-Za-z]+", given or "")
    return "".join(part[0].upper() for part in parts)


def format_reference(number, record):
    authors = record.get("authors", [])
    names = [f"{author.get('family', '').strip()} {initials(author.get('given', ''))}".strip() for author in authors]
    if len(names) > 6:
        author_text = ", ".join(names[:6]) + ", et al."
    else:
        author_text = ", ".join(names) + "."
    title = clean_markup(record.get("title", ""))
    journal = clean_markup(record.get("container_title") or record.get("short_container_title") or "")
    year = record.get("year") or ""
    volume = record.get("volume") or ""
    issue = record.get("issue") or ""
    page = record.get("page") or ""
    citation = f"{number}. {author_text} {title}. {journal}. {year}"
    if volume:
        citation += f";{volume}"
        if issue:
            citation += f"({issue})"
    if page:
        citation += f":{page}"
    citation += f". doi:{record['doi']}"
    return re.sub(r"\s+", " ", citation)


def reference_list():
    records = json.loads(REFERENCE_JSON.read_text(encoding="utf-8"))
    by_doi = {record["doi"].lower(): record for record in records}
    order = [
        "10.1183/09031936.00185114",
        "10.1164/rccm.202202-0399ST",
        "10.1016/s0140-6736(11)60052-4",
        "10.1146/annurev-pathol-042320-030240",
        "10.1165/rcmb.2017-0340OC",
        "10.1152/ajplung.00189.2024",
        "10.26508/lsa.202402805",
        "10.1056/NEJMoa1013660",
        "10.1016/s2213-2600(17)30387-9",
        "10.1038/ng.2609",
        "10.1038/s41586-022-05473-8",
        "10.7554/eLife.34408",
        "10.1001/jama.2021.18236",
        "10.1038/s41588-018-0099-7",
        "10.1093/ije/dyv080",
        "10.1002/gepi.21965",
        "10.1016/j.bbadis.2024.167572",
        "10.1126/sciadv.aba1983",
        "10.1038/s41588-024-01819-2",
        "10.1164/rccm.202306-0979OC",
        "10.1089/omi.2011.0118",
        "10.1093/bioinformatics/btp616",
        "10.1093/nar/gkv007",
        "10.1186/gb-2014-15-2-r29",
        "10.18637/jss.v036.i03",
        "10.1038/s41467-021-25960-2",
        None,
        "10.1111/j.2517-6161.1995.tb02031.x",
        "10.1093/nar/gks1193",
        "10.1016/j.cell.2021.04.048",
        "10.25080/Majora-92bf1922-00a",
        "10.1056/NEJMra1705751",
        "10.1172/jci74942",
        "10.1126/sciadv.aba1972",
        "10.1038/s41467-020-15647-5",
        "10.1161/CIRCULATIONAHA.120.052318",
        "10.1038/s41590-018-0276-y",
        "10.1126/scitranslmed.3005964",
        "10.1016/S2213-2600(17)30349-1",
    ]
    output = []
    for number, doi in enumerate(order, start=1):
        if doi is None:
            output.append(f"{number}. Lawson CL, Hanson RJ. Solving Least Squares Problems. Englewood Cliffs, NJ: Prentice-Hall; 1974.")
        else:
            output.append(format_reference(number, by_doi[doi.lower()]))
    return output


def count_words(texts):
    return sum(len(re.findall(r"\b[\w'-]+\b", text)) for text in texts)


def build_manuscript():
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(28)
    title.paragraph_format.space_after = Pt(15)
    run = title.add_run("Cell-Contextual Metabolic Remodeling in Idiopathic Pulmonary Fibrosis: An Integrative Genetic and Transcriptomic Study")
    set_font(run, size=17, bold=True, color=RGBColor(22, 41, 55))

    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_data = [
        ("Sutong Chen, MD", "1,*"),
        ("Yuhua Fang, BSc", "2,*"),
        ("Xinyi Kong, MSc", "3"),
        ("Juan Wu, BSc", "1"),
        ("Liu Yang, MD", "1"),
        ("Wei Xie, MSc", "1"),
        ("Zanyan Wan, MSc", "1"),
        ("Jie Li, MD", "1"),
    ]
    for index, (name, affiliation) in enumerate(author_data):
        if index:
            set_font(authors.add_run("; "), size=11.2)
        set_font(authors.add_run(name), size=11.2)
        superscript = authors.add_run(affiliation)
        set_font(superscript, size=8.2)
        superscript.font.superscript = True
    authors.paragraph_format.space_after = Pt(10)

    affiliations = [
        "1 Department of Tuberculosis, Jiangxi Chest Hospital, The Third Affiliated Hospital of Nanchang Medical College, Nanchang, Jiangxi 330006, China",
        "2 Huankui Academy, Nanchang University, Nanchang, Jiangxi 330031, China",
        "3 Department of Cardiac Intervention, The Second Affiliated Hospital of Nanchang University, Nanchang, Jiangxi 330006, China",
    ]
    for text in affiliations:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(text)
        set_font(run, size=9.5)
    equal = doc.add_paragraph()
    equal.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(equal.add_run("* These authors contributed equally to this work."), size=9.5, italic=True)

    corr = doc.add_paragraph()
    corr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(corr.add_run("Correspondence: "), size=9.5, bold=True)
    set_font(corr.add_run("Jie Li, MD; Department of Tuberculosis, Jiangxi Chest Hospital, Nanchang, Jiangxi 330006, China; lijie_3608@126.com; +86 152-7083-1364"), size=9.5)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(12)
    set_font(meta.add_run("Article type: Research Article | Running title: Cell-contextual metabolism in IPF"), size=9.5, color=MUTED)
    counts_placeholder = doc.add_paragraph()
    counts_placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
    counts_run = counts_placeholder.add_run("WORD_COUNTS_PLACEHOLDER")
    set_font(counts_run, size=9.5, color=MUTED)

    doc.add_page_break()

    doc.add_heading("Abstract", level=1)
    abstract_parts = []
    abstract_data = [
        ("Background", "Metabolic dysregulation is implicated in idiopathic pulmonary fibrosis (IPF), but whole-lung transcriptional changes may reflect both cell-intrinsic regulation and altered cellular composition. We sought to separate reproducible metabolic signals from broad fibrotic remodeling and to resolve their lung-cell context."),
        ("Methods", "We integrated a nominal genetic screen with purified alveolar epithelial type II (AT2) RNA sequencing, three independent whole-lung cohorts, lung single-cell RNA sequencing, and composition-sensitivity analyses. Cross-platform whole-lung effects were standardized as Hedges' g and pooled with Hartung-Knapp random-effects models. Expression-matched permutations assessed candidate-set specificity. GSE136831 was used for donor-level FASN localization, a six-gene candidate metabolic module, and marker-reference sensitivity analyses."),
        ("Results", "The genetic archive contained 147 nominal associations but lacked exposure-level provenance and a complete testing universe; this layer was therefore treated as exploratory. In GSE245965, 806 genes were upregulated and 800 were downregulated at BH-FDR < 0.05 and |log2 fold change| > 1. GSEA showed positive epithelial-remodeling programs and negative cholesterol, fatty-acid, surfactant, unfolded-protein-response, phospholipid, and oxidative-phosphorylation programs. FASN decreased directionally in all three whole-lung cohorts, but the standardized pooled effect did not reach conventional significance (Hedges' g = -1.05, 95% CI -2.13 to 0.03; P = 0.052; candidate-set FDR = 0.838; I2 = 50.1%). No candidate passed meta-FDR < 0.10, and candidate recurrence did not exceed matched-background expectations. Lung single-cell analysis localized FASN to alveolar epithelial compartments and identified lower expression in IPF AT1 cells (FDR = 0.029), whereas AT2 differences were not significant. A post hoc candidate metabolic module increased in endothelial cells (FDR = 6.8 x 10^-5). Marker-score and NNLS adjustment materially altered the whole-lung FASN effect, with wide confidence intervals in smaller cohorts."),
        ("Conclusions", "A directionally recurrent whole-lung FASN signal is not sufficient to establish a causal target or a cell-intrinsic decrease. Cell-resolved and composition-sensitive analyses reframe FASN as an example of context-dependent metabolic remodeling in IPF and emphasize the need for transparent genetic annotation, standardized cross-platform synthesis, and cell-specific validation."),
    ]
    for label, text in abstract_data:
        add_labeled_paragraph(doc, label + ":", text)
        abstract_parts.append(text)
    keywords = doc.add_paragraph()
    set_font(keywords.add_run("Keywords: "), bold=True)
    set_font(keywords.add_run("idiopathic pulmonary fibrosis; FASN; metabolism; single-cell RNA sequencing; tissue composition; transcriptomic meta-analysis"))

    main_texts = []
    doc.add_heading("Introduction", level=1)
    intro = [
        "Idiopathic pulmonary fibrosis (IPF) is a progressive fibrosing interstitial lung disease with a substantial mortality burden and limited capacity for tissue repair [1-4]. Recurrent alveolar epithelial injury, failed regeneration, fibroblast and myofibroblast activation, vascular remodeling, immune recruitment, and extracellular-matrix accumulation coexist within the same diseased lung. Antifibrotic therapies slow functional decline but do not reverse established fibrosis, making the biological interpretation of molecular targets an important translational problem [2,32].",
        "Metabolic remodeling is one plausible link between epithelial stress and fibrotic persistence. Lipid synthesis supports membrane turnover, surfactant production, and resolution of endoplasmic-reticulum stress, whereas fibroblast fatty-acid and mTOR-dependent programs can support activated phenotypes [5-7]. These observations create a central interpretive tension: a pathway that is protective in an epithelial compartment may be profibrotic in a stromal compartment. A bulk-lung expression change cannot distinguish cell-intrinsic regulation from loss or expansion of the cells that express the gene.",
        "Genetic association studies have identified IPF susceptibility loci related to mucin biology, epithelial defense, telomere maintenance, and host response [8-10]. Mendelian randomization (MR) can reduce some forms of confounding and reverse causation, but its interpretation depends on transparent exposure definitions, instrument provenance, directionality, pleiotropy assessment, and correction across the complete testing universe [11-16]. Whole-lung transcriptomic replication can identify recurring directions but remains vulnerable to platform differences and altered tissue composition. Single-cell data can localize signals, yet donor-level inference is required to avoid pseudoreplication [18,26].",
        "We hypothesized that metabolism-associated IPF signals could recur across lung cohorts while their apparent direction in whole tissue would be modified by lung cellular composition. We therefore treated the available genetic results as a nominal screening layer, reanalyzed purified AT2 RNA sequencing with actual enrichment statistics, synthesized three whole-lung cohorts using standardized effects, calibrated candidate recurrence against expression-matched genes, and used lung single-cell and composition-sensitivity analyses to test how cellular context altered the interpretation of FASN. PBMC single-cell data were retained only as a supplementary systemic layer [17-20].",
    ]
    for text in intro:
        add_body(doc, text)
        main_texts.append(text)

    doc.add_heading("Methods", level=1)
    doc.add_heading("Study design and data sources", level=2)
    methods = [
        "This study integrated public genetic and transcriptomic resources without generating new human or animal samples. Dataset accession, sample composition, platform, analytical role, and interpretive boundary are summarized in Table 1. Transcriptomic datasets were accessed through the NCBI Gene Expression Omnibus (GEO) [29]. The IPF genetic outcome was FinnGen Release 12 endpoint IPF, which included 2,700 cases and 494,583 controls according to the public release manifest [11].",
    ]
    for text in methods:
        add_body(doc, text)
        main_texts.append(text)
    add_dataset_table(doc)

    doc.add_heading("Nominal genetic screening", level=2)
    text = "The available project archive contained inverse-variance-weighted estimates for 147 nominal metabolism-associated genetic exposures and selected heterogeneity and pleiotropy diagnostics, but it did not contain exposure-level GWAS/QTL provenance, units, ancestry, sample size, cis/trans classification, harmonization logs, or the complete testing universe. Consequently, this layer was reclassified as a retrospective nominal genetic screen rather than causal MR discovery. The archived workflow used genome-wide-significant instruments where available, linkage-disequilibrium clumping at r2 < 0.001 within 10,000 kb, exclusion of instruments with F statistics below 10, inverse-variance weighting as the primary estimator, and weighted-median, MR-Egger, MR-PRESSO, heterogeneity, and leave-one-out checks when instrument counts permitted [12-16]. The 147 associations at nominal P < 0.05 defined a broad candidate space. The previous practice of applying Benjamini-Hochberg adjustment only after selecting these 147 associations was not used as inferential FDR because post-selection correction does not represent the complete testing universe [28]."
    add_body(doc, text)
    main_texts.append(text)

    doc.add_heading("Purified AT2 transcriptomic analysis", level=2)
    text = "GSE245965 was analyzed as purified AT2-cell RNA sequencing, with two IPF and three normal samples; ATAC-seq and basal-cell samples were excluded [17]. Raw counts were filtered with edgeR, normalized by trimmed mean of M values, transformed with limma-voom, and modeled with group as the design factor [22-24]. Differential-expression candidates were defined by BH-FDR < 0.05 and |log2 fold change| > 1. The moderated t statistic was used to rank 18,026 measured genes for preranked GSEA. Hallmark and Reactome gene sets were obtained through msigdbr; fgsea multilevel estimation used gene-set sizes of 10-500 genes and BH correction across tested pathways. Figure 2 displays actual NES, FDR, and gene-set size rather than ordinal summary scores [21,28]."
    add_body(doc, text)
    main_texts.append(text)

    doc.add_heading("Whole-lung standardization and meta-analysis", level=2)
    text = "Independent whole-lung cohorts comprised GSE150910 (103 IPF and 103 controls; RNA sequencing), GSE110147 (22 IPF and 11 controls; microarray), and GSE24206 (17 IPF and 6 healthy controls; microarray). RNA-seq counts were normalized with edgeR and voom, and microarrays were analyzed on their normalized expression scales with probe-to-gene collapse based on highest interquartile variability [22-24]. To avoid pooling incomparable platform-specific log fold changes, each gene was converted to Hedges' g using the pooled within-cohort standard deviation and small-sample correction. Sampling variances were calculated from group sizes and the standardized effect. Candidate effects were pooled using restricted-maximum-likelihood random-effects models with Hartung-Knapp inference; heterogeneity was summarized by I2 [25]. BH-FDR was calculated across all candidate genes with estimable pooled effects. The earlier raw-logFC meta-analysis was retained only as a historical sensitivity output and was not used for the revised primary conclusion."
    add_body(doc, text)
    main_texts.append(text)

    doc.add_heading("Matched-background calibration", level=2)
    text = "To test whether candidate recurrence exceeded broad IPF transcriptional remodeling, expression-matched permutations were performed separately in each whole-lung cohort and for support in at least two cohorts. Candidate genes were matched to measured background genes by mean-expression bins; 1,000 random sets were sampled, and the observed number of candidate-level FDR < 0.10 genes was compared with the null distribution. This calibration evaluates gene-set specificity rather than the validity of individual effects."
    add_body(doc, text)
    main_texts.append(text)

    doc.add_heading("Lung single-cell localization", level=2)
    text = "GSE136831 whole-lung dissociation single-cell RNA sequencing was used after excluding COPD cells and multiplets [18]. The analyzed archive comprised 28 control and 32 IPF donors and 230,550 cells after study filters. Candidate expression was summarized across AT2, AT1, other epithelial, fibroblast, myofibroblast, macrophage, monocyte, endothelial, T/NK, and B/plasma compartments; aberrant basaloid cells were used for localization but not balanced case-control testing. Raw target-gene counts were aggregated by donor and cell type, then converted to log1p counts per 10,000 UMIs for descriptive donor-level comparisons. Cell types required at least five cells per donor stratum. Wilcoxon rank-sum tests compared control and IPF donors, with BH correction across tested cell types. This donor-level aggregation avoids treating cells as independent biological replicates, but it remains an exploratory alternative to full raw-count pseudobulk models with covariates [26]."
    add_body(doc, text)
    main_texts.append(text)

    doc.add_heading("Candidate metabolic module and composition sensitivity", level=2)
    text = "A six-gene set (FASN, ACACB, LPCAT2, GLO1, PGM2L1, and GGCX) was retained as a post hoc candidate metabolic module, not as a canonical FASN or lipid pathway. Its donor-level mean log1p CP10K score was tested across lung cell types with the same donor and FDR procedures. Whole-lung composition sensitivity was evaluated first with z-scored marker averages for AT2, fibroblast, myofibroblast, macrophage, and endothelial compartments and then with non-negative least squares (NNLS) estimates derived from a GSE136831 marker reference [27]. The NNLS models adjusted FASN expression for AT2, AT1, fibroblast, myofibroblast, macrophage, and endothelial estimates. These covariates are approximate reference-dependent components rather than measured fractions; attenuation was interpreted as evidence of composition sensitivity, not positive support for FASN."
    add_body(doc, text)
    main_texts.append(text)

    doc.add_heading("Supplementary PBMC analysis", level=2)
    text = "GSE233844 contains 38 PBMC samples from 13 controls, 13 patients with stable IPF, and 12 with progressive IPF, yielding 149,564 cells [20]. The archived pathway summaries were retained in Supplementary Results as a systemic disease layer. They were not used for lung localization, target ranking, or prognostic claims. Seurat was used for major-lineage annotation in the original pipeline [30]."
    add_body(doc, text)
    main_texts.append(text)

    doc.add_heading("Statistical analysis and reproducibility", level=2)
    text = "All tests were two-sided. BH-FDR was applied across the complete pathway, cell-type, or candidate set relevant to each analysis [28]. Effect estimates are reported with 95% confidence intervals. R was used for expression modeling, enrichment, and meta-analysis; Python was used for structured data handling, figure assembly, and document generation [21-25,31]. Revised figure scripts and source-data tables accompany the manuscript."
    add_body(doc, text)
    main_texts.append(text)

    doc.add_heading("Results", level=1)
    doc.add_heading("A nominal genetic screen defined a broad but incompletely annotated candidate space", level=2)
    text = "The archived screen contained 147 metabolism-associated nominal results at inverse-variance-weighted P < 0.05. Instrument counts ranged from 3 to 112 (median 20; interquartile range 10-36). Selected MR-Egger-intercept and MR-PRESSO global tests did not show strong evidence of directional pleiotropy, but these diagnostics were available for only a subset and cannot compensate for absent exposure provenance or full-universe multiple-testing control. We therefore use the genetic screen only to define candidate context. Figure 1 ranks Panel B strictly by archived nominal P value, removes the previous cross-layer selection loop, and replaces the subjective metabolic-axis panel with the observed instrument-count distribution."
    add_body(doc, text)
    main_texts.append(text)
    add_figure(
        doc,
        1,
        "Nominal genetic screening and reporting calibration. (A) Archived workflow from FinnGen IPF outcome to cell-context refinement. (B) Twelve lowest-P nominal inverse-variance-weighted associations, selected only by the genetic-screen P value. (C) Available heterogeneity and pleiotropy diagnostics for selected candidates. (D) Number of SNP instruments across the 147 nominal associations. Instrument count does not establish instrument validity; exposure provenance and full testing-universe FDR remain required. IVW, inverse-variance weighted; MR, Mendelian randomization; IPF, idiopathic pulmonary fibrosis.",
    )

    doc.add_heading("Purified AT2 cells showed opposing structural and metabolic programs", level=2)
    text = "The revised limma-voom analysis of GSE245965 measured 18,026 genes and identified 806 upregulated and 800 downregulated genes at BH-FDR < 0.05 and |log2 fold change| > 1. Positive GSEA signals included epithelial-mesenchymal transition (NES 2.32; FDR 1.75 x 10^-9), extracellular-matrix organization (NES 2.15; FDR 6.97 x 10^-9), keratinization, and cornified-envelope formation. In contrast, cholesterol biosynthesis (NES -2.61; FDR 2.71 x 10^-7), Hallmark fatty-acid metabolism (NES -2.17; FDR 7.30 x 10^-7), surfactant metabolism (NES -2.38; FDR 5.08 x 10^-5), unfolded-protein response (NES -1.94; FDR 0.0013), phospholipid metabolism (NES -1.67; FDR 0.0036), and oxidative phosphorylation (NES -1.59; FDR 0.0076) were negatively enriched. These results replace the previous ordinal enrichment scores with reproducible statistics and support an epithelial state combining structural remodeling with reduced lipid and energy programs."
    add_body(doc, text)
    main_texts.append(text)
    add_figure(
        doc,
        2,
        "Purified AT2-cell differential expression and actual pathway enrichment in GSE245965. (A) Volcano plot using BH-FDR < 0.05 and |log2 fold change| > 1; two IPF and three normal samples were analyzed. (B) Positively enriched structural and remodeling programs. (C) Negatively enriched metabolic and stress-response programs. Horizontal position is the GSEA normalized enrichment score (NES), point size is gene-set size, and color is -log10 BH-FDR. No ordinal or manually assigned pathway score is shown. AT2, alveolar epithelial type II cell; GSEA, gene-set enrichment analysis.",
    )

    doc.add_heading("Standardized whole-lung synthesis showed directional FASN recurrence without FDR-level discovery", level=2)
    text = "FASN was lower in IPF in each whole-lung cohort after platform-specific normalization: GSE150910 Hedges' g -0.75 (95% CI -1.03 to -0.47), GSE110147 -1.58 (95% CI -2.41 to -0.76), and GSE24206 -1.18 (95% CI -2.17 to -0.18). With only three cohorts, Hartung-Knapp random-effects inference widened the pooled interval: Hedges' g -1.05 (95% CI -2.13 to 0.03), P = 0.052, I2 = 50.1%. After BH correction across the estimable candidate set, FASN FDR was 0.838 and no candidate reached meta-FDR < 0.10. Thus, FASN showed directionally recurrent cohort-level decreases but not a statistically secure cross-platform discovery under the revised primary model."
    add_body(doc, text)
    main_texts.append(text)
    text = "Matched-background permutations further limited gene-set claims. Empirical P values were 0.627 for GSE150910, 0.966 for GSE110147, 0.294 for GSE24206, and 0.708 for support in at least two cohorts. The nominal genetic candidate set was therefore not globally enriched beyond expression-matched background genes. The defensible inference is candidate-specific contextualization, not enrichment of a metabolic MR gene set."
    add_body(doc, text)
    main_texts.append(text)
    add_figure(
        doc,
        3,
        "Standardized whole-lung analysis and negative calibration. (A) Cohort-specific FASN Hedges' g values and Hartung-Knapp random-effects summary. (B) Candidate meta-analysis P values compared with the BH-FDR 0.10 boundary; no candidate crossed the boundary. Point color represents I2. (C) Empirical P values from expression-matched background permutations for each cohort and for support in at least two cohorts. CI, confidence interval; FDR, false-discovery rate; IPF, idiopathic pulmonary fibrosis.",
    )

    doc.add_heading("Cellular context altered the interpretation of whole-lung FASN", level=2)
    text = "FASN was most prominent in alveolar epithelial compartments in GSE136831. Donor-level summaries showed lower FASN in IPF AT1 cells (median difference -0.227 log1p CP10K; FDR 0.029; 15 control and 9 IPF donors), whereas the AT2 difference was smaller and not significant (median difference -0.064; FDR 0.953; 17 control and 15 IPF donors). Other epithelial and fibroblast compartments showed negative directions without FDR-level support. These results do not establish a universal epithelial decrease, but they identify AT1 as the only tested compartment with donor-level FDR support."
    add_body(doc, text)
    main_texts.append(text)
    text = "The six-gene candidate metabolic module did not show a pan-cell decrease. Its strongest difference was an endothelial increase (median difference 0.283 log1p CP10K; FDR 6.8 x 10^-5; 25 control and 31 IPF donors), whereas epithelial and stromal effects were not significant after correction. Because this is a post hoc candidate module rather than a canonical pathway, the endothelial result is interpreted as a compartment-specific descriptive signal requiring independent validation."
    add_body(doc, text)
    main_texts.append(text)
    text = "Composition adjustments materially changed the whole-lung FASN coefficient. Marker-score adjustment attenuated the negative effect in GSE150910 and GSE110147 but amplified it in GSE24206, demonstrating reference dependence rather than uniform attenuation. NNLS adjustment produced non-significant effects in all cohorts: -0.23 (95% CI -0.58 to 0.11; P = 0.189), 0.26 (95% CI -1.95 to 2.48; P = 0.811), and -0.66 (95% CI -1.92 to 0.59; P = 0.278), respectively. The extremely wide GSE110147 interval indicates model instability in a small cohort. Together, these analyses show that the tissue-level FASN direction is composition-sensitive and cannot be assigned to a single cell-intrinsic mechanism."
    add_body(doc, text)
    main_texts.append(text)
    add_figure(
        doc,
        4,
        "Lung single-cell localization and whole-lung composition sensitivity. (A) Expression of eight preselected metabolism-associated candidates across major GSE136831 lung cell types. (B) Donor-level FASN summaries; facet strips report control and IPF donor numbers and BH-FDR. (C) Donor-level post hoc candidate metabolic module comprising FASN, ACACB, LPCAT2, GLO1, PGM2L1, and GGCX; this is not a canonical FASN pathway. (D) Whole-lung FASN coefficients before adjustment, after marker-score adjustment, and after NNLS marker-reference adjustment. NNLS estimates are sensitivity covariates rather than measured cell fractions. CP10K, counts per 10,000 UMIs; NNLS, non-negative least squares.",
    )

    doc.add_heading("Discussion", level=1)
    discussion = [
        "The principal finding is that a directionally recurrent whole-lung FASN decrease becomes substantially less certain when platform differences, multiple testing, cellular localization, and tissue composition are considered together. FASN decreased in each individual bulk cohort, but the conservative standardized pooled interval crossed zero and no candidate passed meta-FDR < 0.10. The candidate set also failed matched-background enrichment. These negative calibrations are not peripheral caveats; they define the evidential ceiling of the study.",
        "Cellular resolution changed the biological interpretation of the bulk signal. FASN localized mainly to alveolar epithelial compartments, and the only donor-level FDR-supported difference occurred in AT1 rather than AT2 cells. Composition adjustment changed the whole-lung coefficient in a dataset- and reference-dependent manner. This pattern is compatible with altered abundance of alveolar and structural compartments contributing to the tissue average, while leaving open the possibility of cell-intrinsic regulation in selected states. It is not compatible with describing FASN as a uniformly downregulated therapeutic target.",
        "The AT2 analysis nevertheless provides a coherent metabolic context. Positive epithelial-remodeling and extracellular-matrix programs coexisted with negative cholesterol, fatty-acid, surfactant, phospholipid, unfolded-protein-response, and oxidative-phosphorylation programs. Lipid synthesis can support resolution of epithelial endoplasmic-reticulum stress [5], and impaired mitochondrial quality control can promote fibrotic susceptibility [33]. Conversely, FASN inhibition has reduced fibroblast activation in experimental systems [7]. These observations reinforce a cell-dependent model in which restoring epithelial metabolic competence and inhibiting stromal metabolic activation are not interchangeable interventions.",
        "The endothelial increase in the six-gene candidate module further argues against a single pan-lung direction. The lung endothelium is metabolically active and participates in barrier function, trafficking, and repair [36]. Single-cell atlases have identified abnormal epithelial, stromal, endothelial, and macrophage states in fibrotic lungs [18,34-37]. Our module, however, was assembled from prioritized candidates and includes genes outside canonical fatty-acid synthesis. Its endothelial association should therefore motivate pathway-defined and leave-one-gene-out validation rather than support a new named axis.",
        "The genetic layer requires the greatest caution. Instrument-count summaries and selected pleiotropy diagnostics cannot replace exposure-level QTL/GWAS identifiers, units, ancestry, sample overlap, harmonization decisions, Steiger directionality, and correction across all tested exposures [13-16]. Until those materials are restored and targeted colocalization is performed, the 147 associations are best regarded as an exploratory screening space. The PBMC analysis is similarly supplementary: it can describe systemic immune context, but without a revised donor-level model and independent validation it should not drive target or prognosis claims [20,38,39].",
        "Several limitations remain. GSE245965 contains only five RNA-seq samples, so AT2 enrichment estimates are exploratory. The standardized whole-lung meta-analysis includes only three cohorts, making Hartung-Knapp intervals appropriately wide. GSE136831 preserves cell identity but not spatial niche, and the donor summaries use log-normalized target counts rather than a fully covariate-adjusted raw-count pseudobulk model. Marker-score and NNLS deconvolution are reference-dependent and underpowered in the smaller cohorts; condition-number, variance-inflation, bootstrap, leave-one-marker-out, and alternative-reference analyses remain necessary. No wet-laboratory perturbation was performed. In conclusion, the current data support cell-context-dependent metabolic remodeling, with FASN as a useful example of why bulk-tissue direction should not be equated with a causal, unidirectional target.",
    ]
    for text in discussion:
        add_body(doc, text)
        main_texts.append(text)

    doc.add_heading("Declarations", level=1)
    declarations = [
        ("Author contributions", "A final CRediT statement was not available in the source manuscript and must be confirmed by all authors before submission."),
        ("Funding", "Funding information was not available in the source manuscript and must be completed before submission."),
        ("Competing interests", "Author-specific competing-interest declarations were not available in the source manuscript and must be completed before submission."),
        ("Ethics approval", "This secondary analysis used de-identified public datasets. Ethics approval and informed consent were reported by the original studies; no new participants were enrolled."),
        ("Consent for publication", "Not applicable to this secondary analysis of de-identified public data."),
        ("Data availability", "Public data are available under FinnGen Release 12 endpoint IPF and GEO accessions GSE245965, GSE150910, GSE110147, GSE24206, GSE136831, and GSE233844. Revised figure source data accompany the project package. A permanent public repository and DOI should be added before publication."),
        ("Code availability", "Scripts used for the revised AT2 enrichment analysis, standardized whole-lung meta-analysis, and Figures 1-4 accompany the project package. A versioned public archive should be created before publication."),
        ("Acknowledgements", "No acknowledgement statement was available in the source manuscript."),
        ("Supplementary material", "The revised supplementary workbook contains dataset metadata, actual AT2 GSEA results, cohort-specific and pooled standardized whole-lung effects, background calibration, lung single-cell donor tests, and composition-sensitivity models. The PBMC analysis is retained as supplementary systemic context."),
    ]
    for label, text in declarations:
        add_labeled_paragraph(doc, label + ":", text)

    doc.add_heading("References", level=1)
    for reference in reference_list():
        paragraph = doc.add_paragraph(style="Reference")
        run = paragraph.add_run(reference)
        set_font(run, size=9)

    abstract_count = count_words(abstract_parts)
    main_count = count_words(main_texts)
    for paragraph in doc.paragraphs:
        if "WORD_COUNTS_PLACEHOLDER" in paragraph.text:
            paragraph.clear()
            run = paragraph.add_run(
                f"Word counts: Abstract {abstract_count}; main text {main_count} | Tables 1 | Figures 4 | Supplementary material: Yes"
            )
            set_font(run, size=9.5, color=MUTED)

    doc.core_properties.title = "Cell-Contextual Metabolic Remodeling in Idiopathic Pulmonary Fibrosis"
    doc.core_properties.subject = "Revised IPF integrative transcriptomic manuscript"
    doc.core_properties.keywords = "IPF; FASN; metabolism; single-cell RNA sequencing; tissue composition"
    doc.core_properties.author = "Sutong Chen et al."
    doc.save(OUTPUT)
    return abstract_count, main_count


def build_action_items():
    doc = Document()
    configure_document(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Submission Action Items for the Revised IPF Manuscript")
    set_font(run, size=17, bold=True, color=RGBColor(22, 41, 55))
    lead = doc.add_paragraph()
    lead.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = lead.add_run(
        "The scientific narrative and statistical figures have been revised. The following items remain author-controlled or require unavailable source data and must be completed before journal submission."
    )
    set_font(run, size=11)

    items = [
        ("P0", "Restore genetic exposure provenance", "For every tested exposure, provide accession/ID, biological quantity, GWAS/QTL type, ancestry, sample size, unit, cis/trans definition, SNP-selection rules, harmonization decisions, proxies, sample overlap, and instrument-level F statistics."),
        ("P0", "Reconstruct the complete testing universe", "Document the total number of exposures tested and calculate BH-FDR across that complete set. Do not use FDR calculated only after selecting the 147 nominal hits."),
        ("P0", "Confirm author statements", "Complete CRediT contributions, funding, competing interests, acknowledgements, and ORCID iDs. Confirm degrees and official English institutional names."),
        ("P0", "Create public code/data archive", "Deposit the revised scripts, package versions, source-data tables, and figure files in a versioned repository; mint a permanent DOI and update Data/Code Availability."),
        ("P1", "Strengthen genetic inference", "Add Steiger directionality and targeted colocalization for the small set of candidates retained after complete-universe calibration. Use cis instruments where biologically appropriate."),
        ("P1", "Upgrade lung pseudobulk", "Aggregate full raw UMI matrices by donor and cell type, fit edgeR/DESeq2/limma-voom with donor as the biological replicate, report donor counts, and assess >=20 and >=50 cell thresholds with available covariates."),
        ("P1", "Stress-test deconvolution", "Report signature-matrix condition number and VIF; add bootstrap confidence intervals, leave-one-marker-out and leave-one-cell-type-out analyses, and one alternative method/reference atlas."),
        ("P1", "Validate the candidate module", "Replace or complement the six-gene post hoc module with a canonical lipid pathway; perform leave-one-gene-out tests to determine whether the endothelial signal is driven by a single gene."),
        ("P2", "Keep PBMC results supplementary", "Do not restore PBMC pathway summaries to the main Results unless donor-level effect sizes, FDR, and an independent validation cohort are added."),
    ]
    table = doc.add_table(rows=1, cols=4)
    headers = ["Priority", "Item", "Required action", "Status"]
    for i, value in enumerate(headers):
        table.rows[0].cells[i].text = value
        set_cell_shading(table.rows[0].cells[i], HEADER_FILL)
        for paragraph in table.rows[0].cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_font(run, size=9, bold=True)
    set_repeat_table_header(table.rows[0])
    for row_index, (priority, item, action, status) in enumerate(
        [(priority, item, action, "Open") for priority, item, action in items], start=1
    ):
        cells = table.add_row().cells
        values = [priority, item, action, status]
        for col, value in enumerate(values):
            cells[col].text = value
            cells[col].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if row_index % 2 == 0:
                set_cell_shading(cells[col], "F7F9FB")
            for paragraph in cells[col].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_font(run, size=8.5, bold=(col == 0))
    set_table_geometry(table, [0.55, 1.2, 4.15, 0.55])
    doc.save(ACTION_OUTPUT)


def main():
    for path in [REFERENCE_JSON, *FIGURES.values()]:
        if not path.exists():
            raise FileNotFoundError(path)
    abstract_count, main_count = build_manuscript()
    build_action_items()
    print(f"Wrote {OUTPUT}")
    print(f"Wrote {ACTION_OUTPUT}")
    print(f"Abstract words: {abstract_count}; main-text words: {main_count}")


if __name__ == "__main__":
    main()
