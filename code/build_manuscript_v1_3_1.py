from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

import document_helpers as base


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build the IPF source-reanalysis v1.3.1 manuscript.")
    parser.add_argument("--analysis-root", required=True, type=Path)
    return parser.parse_args()


ARGS = parse_args()
ANALYSIS = ARGS.analysis_root.resolve()
OUTPUT = ANALYSIS / "Manuscript_v1.3.1_20260712.docx"
ACTION_OUTPUT = ANALYSIS / "Submission_Action_Items_v1.3.1_20260712.docx"

FIGURES = {
    1: ANALYSIS / "Figures_v1.3.1_20260712" / "Main" / "Figure1_study_architecture.png",
    2: ANALYSIS / "Figures_v1.3.1_20260712" / "Main" / "Figure2_AT2_structural_metabolic_divergence.png",
    3: ANALYSIS / "Figures_v1.3.1_20260712" / "Main" / "Figure3_whole_lung_source_synthesis.png",
    4: ANALYSIS / "Figures_v1.3.1_20260712" / "Main" / "Figure4_full_transcriptome_cell_context.png",
}
base.FIGURES = FIGURES


def count_words(texts: list[str]) -> int:
    return sum(len(text.split()) for text in texts)


def add_dataset_table(doc: Document) -> None:
    base.add_table_caption(doc, 1, "Source datasets, sample composition, and analytical roles.")
    headers = ["Dataset", "Modality", "Source input", "Samples", "Primary role", "Boundary"]
    rows = [
        ["GSE245965", "Purified AT2 RNA-seq", "Raw counts", "2 IPF; 3 controls", "AT2 differential expression and pathway analysis", "Very small cohort"],
        ["GSE150910", "Whole-lung RNA-seq", "Raw counts", "103 IPF; 103 controls", "Standardized synthesis and composition sensitivity", "Bulk tissue"],
        ["GSE110147", "Whole-lung microarray", "48 raw CEL; RMA", "22 IPF; 11 controls", "Standardized synthesis and composition sensitivity", "Adjusted models have high collinearity"],
        ["GSE24206", "Whole-lung microarray", "23 raw CEL; RMA", "17 IPF; 6 controls", "Standardized synthesis and composition sensitivity", "Small cohort"],
        ["GSE136831", "Lung scRNA-seq", "Full raw UMI matrix", "32 IPF; 28 controls; 232,056 cells", "Full-transcriptome donor pseudobulk and localization", "Donor-specific libraries nested in disease"],
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    base.set_repeat_table_header(table.rows[0])
    for index, value in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = value
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        base.set_cell_shading(cell, base.HEADER_FILL)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                base.set_font(run, size=7.8, bold=True, color=RGBColor(25, 45, 60))
    for row_index, values in enumerate(rows, start=1):
        cells = table.add_row().cells
        for column_index, value in enumerate(values):
            cell = cells[column_index]
            cell.text = value
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if row_index % 2 == 0:
                base.set_cell_shading(cell, "F7F9FB")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    base.set_font(run, size=7.3)
    base.set_table_geometry(table, [0.75, 0.9, 1.0, 1.05, 1.5, 1.3])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def references() -> list[str]:
    return [
        "1. Hutchinson J, Fogarty A, Hubbard R, McKeever T. Global incidence and mortality of idiopathic pulmonary fibrosis: a systematic review. European Respiratory Journal. 2015;46(3):795-806. doi:10.1183/09031936.00185114",
        "2. Raghu G, Remy-Jardin M, Richeldi L, Thomson CC, Inoue Y, Johkoh T, et al. Idiopathic Pulmonary Fibrosis (an Update) and Progressive Pulmonary Fibrosis in Adults: An Official ATS/ERS/JRS/ALAT Clinical Practice Guideline. American Journal of Respiratory and Critical Care Medicine. 2022;205(9):e18-e47. doi:10.1164/rccm.202202-0399ST",
        "3. King TE, Pardo A, Selman M. Idiopathic pulmonary fibrosis. The Lancet. 2011;378(9807):1949-1961. doi:10.1016/S0140-6736(11)60052-4",
        "4. Moss BJ, Ryter SW, Rosas IO. Pathogenic Mechanisms Underlying Idiopathic Pulmonary Fibrosis. Annual Review of Pathology. 2022;17:515-546. doi:10.1146/annurev-pathol-042320-030240",
        "5. Romero F, Hong X, Shah D, Kallen CB, Rosas I, Guo Z, et al. Lipid Synthesis Is Required to Resolve Endoplasmic Reticulum Stress and Limit Fibrotic Responses in the Lung. American Journal of Respiratory Cell and Molecular Biology. 2018;59(2):225-236. doi:10.1165/rcmb.2017-0340OC",
        "6. Shin KWD, Atalay MV, Cetin-Atalay R, O'Leary EM, Glass ME, Szafran JCH, et al. mTOR signaling regulates multiple metabolic pathways in human lung fibroblasts after TGF-beta and in pulmonary fibrosis. American Journal of Physiology-Lung Cellular and Molecular Physiology. 2025;328(2):L215-L228. doi:10.1152/ajplung.00189.2024",
        "7. Lian H, Zhang Y, Zhu Z, Wan R, Wang Z, Yang K, et al. Fatty acid synthase inhibition alleviates lung fibrosis via beta-catenin signal in fibroblasts. Life Science Alliance. 2025;8(2):e202402805. doi:10.26508/lsa.202402805",
        "8. St. Pierre L, Berhan A, Sung EK, Alvarez JR, Wang H, Ji Y, et al. Integrated multiomic analysis identifies TRIP13 as a mediator of alveolar epithelial type II cell dysfunction in idiopathic pulmonary fibrosis. Biochimica et Biophysica Acta-Molecular Basis of Disease. 2025;1871(3):167572. doi:10.1016/j.bbadis.2024.167572",
        "9. Adams TS, Schupp JC, Poli S, Ayaub EA, Neumark N, Ahangari F, et al. Single-cell RNA-seq reveals ectopic and aberrant lung-resident cell populations in idiopathic pulmonary fibrosis. Science Advances. 2020;6(28):eaba1983. doi:10.1126/sciadv.aba1983",
        "10. Franzen L, Olsson Lindvall M, Huhn M, Ptasinski V, Setyo L, Keith BP, et al. Mapping spatially resolved transcriptomes in human and mouse pulmonary fibrosis. Nature Genetics. 2024;56(8):1725-1736. doi:10.1038/s41588-024-01819-2",
        "11. Yu G, Wang LG, Han Y, He QY. clusterProfiler: an R Package for Comparing Biological Themes Among Gene Clusters. OMICS. 2012;16(5):284-287. doi:10.1089/omi.2011.0118",
        "12. Robinson MD, McCarthy DJ, Smyth GK. edgeR: a Bioconductor package for differential expression analysis of digital gene expression data. Bioinformatics. 2010;26(1):139-140. doi:10.1093/bioinformatics/btp616",
        "13. Ritchie ME, Phipson B, Wu D, Hu Y, Law CW, Shi W, et al. limma powers differential expression analyses for RNA-sequencing and microarray studies. Nucleic Acids Research. 2015;43(7):e47. doi:10.1093/nar/gkv007",
        "14. Law CW, Chen Y, Shi W, Smyth GK. voom: precision weights unlock linear model analysis tools for RNA-seq read counts. Genome Biology. 2014;15(2):R29. doi:10.1186/gb-2014-15-2-r29",
        "15. Viechtbauer W. Conducting Meta-Analyses in R with the metafor Package. Journal of Statistical Software. 2010;36(3):1-48. doi:10.18637/jss.v036.i03",
        "16. Squair JW, Gautier M, Kathe C, Anderson MA, James ND, Hutson TH, et al. Confronting false discoveries in single-cell differential expression. Nature Communications. 2021;12:5692. doi:10.1038/s41467-021-25960-2",
        "17. Lawson CL, Hanson RJ. Solving Least Squares Problems. Englewood Cliffs, NJ: Prentice-Hall; 1974.",
        "18. Benjamini Y, Hochberg Y. Controlling the False Discovery Rate: A Practical and Powerful Approach to Multiple Testing. Journal of the Royal Statistical Society Series B. 1995;57(1):289-300. doi:10.1111/j.2517-6161.1995.tb02031.x",
        "19. Barrett T, Wilhite SE, Ledoux P, Evangelista C, Kim IF, Tomashevsky M, et al. NCBI GEO: archive for functional genomics data sets-update. Nucleic Acids Research. 2013;41(D1):D991-D995. doi:10.1093/nar/gks1193",
        "20. McKinney W. Data Structures for Statistical Computing in Python. Proceedings of the Python in Science Conference. 2010:56-61. doi:10.25080/Majora-92bf1922-00a",
        "21. Lederer DJ, Martinez FJ. Idiopathic Pulmonary Fibrosis. New England Journal of Medicine. 2018;378(19):1811-1823. doi:10.1056/NEJMra1705751",
        "22. Bueno M, Lai YC, Romero Y, Brands J, St. Croix CM, Kamga C, et al. PINK1 deficiency impairs mitochondrial homeostasis and promotes lung fibrosis. Journal of Clinical Investigation. 2015;125(2):521-538. doi:10.1172/JCI74942",
        "23. Habermann AC, Gutierrez AJ, Bui LT, Yahn SL, Winters NI, Calvi CL, et al. Single-cell RNA sequencing reveals profibrotic roles of distinct epithelial and mesenchymal lineages in pulmonary fibrosis. Science Advances. 2020;6(28):eaba1972. doi:10.1126/sciadv.aba1972",
        "24. Tsukui T, Sun KH, Wetter JB, Wilson-Kanamori JR, Hazelwood LA, Henderson NC, et al. Collagen-producing lung cell atlas identifies multiple subsets with distinct localization and relevance to fibrosis. Nature Communications. 2020;11:1920. doi:10.1038/s41467-020-15647-5",
        "25. Schupp JC, Adams TS, Cosme C, Raredon MSB, Yuan Y, Omote N, et al. Integrated Single-Cell Atlas of Endothelial Cells of the Human Lung. Circulation. 2021;144(4):286-302. doi:10.1161/CIRCULATIONAHA.120.052318",
        "26. Aran D, Looney AP, Liu L, Wu E, Fong V, Hsu A, et al. Reference-based analysis of lung single-cell sequencing reveals a transitional profibrotic macrophage. Nature Immunology. 2019;20(2):163-172. doi:10.1038/s41590-018-0276-y",
        "27. Liberzon A, Birger C, Thorvaldsdottir H, Ghandi M, Mesirov JP, Tamayo P. The Molecular Signatures Database Hallmark Gene Set Collection. Cell Systems. 2015;1(6):417-425. doi:10.1016/j.cels.2015.12.004",
        "28. Korotkevich G, Sukhov V, Budin N, Shpak B, Artyomov MN, Sergushichev A. Fast gene set enrichment analysis. bioRxiv. 2021. doi:10.1101/060012",
        "29. Irizarry RA, Hobbs B, Collin F, Beazer-Barclay YD, Antonellis KJ, Scherf U, et al. Exploration, normalization, and summaries of high density oligonucleotide array probe level data. Biostatistics. 2003;4(2):249-264. doi:10.1093/biostatistics/4.2.249",
    ]


def add_front_matter(doc: Document) -> object:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(28)
    title.paragraph_format.space_after = Pt(15)
    base.set_font(
        title.add_run("Cell-Resolved Reanalysis Reframes a Concordant Whole-Lung FASN Signal in Idiopathic Pulmonary Fibrosis"),
        size=17,
        bold=True,
        color=RGBColor(22, 41, 55),
    )
    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_data = [
        ("Sutong Chen, MD", "1,*"), ("Yuhua Fang, BSc", "2,*"), ("Xinyi Kong, MSc", "3"),
        ("Juan Wu, BSc", "1"), ("Liu Yang, MD", "1"), ("Wei Xie, MSc", "1"),
        ("Zanyan Wan, MSc", "1"), ("Jie Li, MD", "1"),
    ]
    for index, (name, affiliation) in enumerate(author_data):
        if index:
            base.set_font(authors.add_run("; "), size=11.2)
        base.set_font(authors.add_run(name), size=11.2)
        superscript = authors.add_run(affiliation)
        base.set_font(superscript, size=8.2)
        superscript.font.superscript = True
    authors.paragraph_format.space_after = Pt(10)
    for text in [
        "1 Department of Tuberculosis, Jiangxi Chest Hospital, The Third Affiliated Hospital of Nanchang Medical College, Nanchang, Jiangxi 330006, China",
        "2 Huankui Academy, Nanchang University, Nanchang, Jiangxi 330031, China",
        "3 Department of Cardiac Intervention, The Second Affiliated Hospital of Nanchang University, Nanchang, Jiangxi 330006, China",
    ]:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(2)
        base.set_font(paragraph.add_run(text), size=9.5)
    equal = doc.add_paragraph()
    equal.alignment = WD_ALIGN_PARAGRAPH.CENTER
    base.set_font(equal.add_run("* These authors contributed equally to this work."), size=9.5, italic=True)
    corr = doc.add_paragraph()
    corr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    base.set_font(corr.add_run("Correspondence: "), size=9.5, bold=True)
    base.set_font(corr.add_run("Jie Li, MD; Department of Tuberculosis, Jiangxi Chest Hospital, Nanchang, Jiangxi 330006, China; lijie_3608@126.com; +86 152-7083-1364"), size=9.5)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(12)
    base.set_font(meta.add_run("Article type: Research Article | Running title: Cell-resolved FASN in IPF"), size=9.5, color=base.MUTED)
    counts_placeholder = doc.add_paragraph()
    counts_placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
    base.set_font(counts_placeholder.add_run("WORD_COUNTS_PLACEHOLDER"), size=9.5, color=base.MUTED)
    doc.add_page_break()
    return counts_placeholder


def add_paragraphs(doc: Document, texts: list[str], word_store: list[str]) -> None:
    for text in texts:
        base.add_body(doc, text)
        word_store.append(text)


def build_manuscript() -> tuple[int, int]:
    doc = Document()
    base.configure_document(doc)
    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.text = "Cell-resolved reanalysis of FASN in IPF"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in header.runs:
            base.set_font(run, size=8.5, italic=True, color=base.MUTED)
    add_front_matter(doc)

    abstract_texts: list[str] = []
    main_texts: list[str] = []
    doc.add_heading("Abstract", level=1)
    abstract_data = [
        ("Background", "A directionally concordant whole-lung signal can reflect intracellular regulation, altered cell abundance, or opposing effects across lung compartments. We integrated source-level purified-cell, whole-lung, and lung single-cell analyses to determine how these factors shaped an IPF-associated FASN signal."),
        ("Methods", "Purified AT2 and whole-lung RNA-seq counts, raw CEL files from two lung cohorts, and the complete GSE136831 raw UMI matrix were reprocessed. Analyses included AT2 sample omission, transparent focal-rule reconstruction, transcriptome-wide Hedges' g meta-analysis with Hartung-Knapp inference, full-transcriptome donor pseudobulk with macrophage subtype separation, low-abundance sensitivity, and composition-model stress tests."),
        ("Results", "Six genes met the AT2/pathway focal rule; FASN had the smallest nominal whole-lung meta-analysis P and was retained as the contextual case. Its AT2 direction remained negative after every sample omission, although FDR ranged from 0.0347 to 0.127. FASN decreased in three whole-lung cohorts, but the pooled interval crossed zero (Hedges' g -1.046; 95% CI -2.100 to 0.009; P=0.0508; transcriptome-wide FDR 0.818). Within GSE136831, ciliated cells showed a threshold-consistent low-abundance negative effect (log2 fold change -1.234; cell-type FDR 8.13x10^-5), whereas alveolar macrophages showed a positive effect (0.490; FDR 0.00686); non-alveolar macrophages were not FDR-supported. Library identifiers were donor-specific and disease-nested, precluding source-library adjustment."),
        ("Conclusions", "FASN showed concordant whole-lung decreases without corrected meta-analytic discovery. Cell-resolved results were compartment- and prevalence-dependent, while source-library confounding could not be separated from disease within the single lung atlas. The data support a contextual case study rather than a uniform mechanism or therapeutic target."),
    ]
    for label, text in abstract_data:
        base.add_labeled_paragraph(doc, label + ":", text)
        abstract_texts.append(text)
    keywords = doc.add_paragraph()
    base.set_font(keywords.add_run("Keywords: "), bold=True)
    base.set_font(keywords.add_run("idiopathic pulmonary fibrosis; FASN; ciliated epithelium; donor pseudobulk; transcriptomic meta-analysis; tissue composition"))

    doc.add_heading("Introduction", level=1)
    add_paragraphs(doc, [
        "Idiopathic pulmonary fibrosis (IPF) is a progressive fibrosing interstitial lung disease driven by recurrent epithelial injury, failed alveolar repair, fibroblast activation, extracellular-matrix deposition, vascular remodeling, and immune-state changes [1-4,25]. Antifibrotic therapies slow functional decline but do not reverse established fibrosis [2,21]. Molecular studies therefore face an important interpretive problem: a reproducible tissue-level gene difference may still be generated by altered cell abundance, a lineage-restricted expression change, or opposing responses among compartments.",
        "Metabolic remodeling is especially sensitive to this problem. Alveolar epithelial type II (AT2) cells require lipid synthesis for membrane turnover, surfactant production, and resolution of endoplasmic-reticulum stress, whereas activated fibroblasts can use fatty-acid and mTOR-linked programs to sustain matrix-producing states [5-7]. Airway-like epithelial remodeling, bronchiolization, and loss of alveolar identity can further shift tissue averages even when expression within a surviving lineage changes in the opposite direction [9,23]. A negative whole-lung average therefore cannot distinguish reduced intracellular expression from lineage loss, expansion of low-expressing populations, or a mixture of mechanisms.",
        "Fatty acid synthase (FASN) provides a tractable case study because it catalyzes de novo fatty-acid synthesis and is measurable across epithelial, stromal, and immune compartments. Experimental observations support context-dependent functions: epithelial lipid synthesis can limit stress-associated fibrotic responses, whereas FASN inhibition can reduce fibroblast activation [5,7]. A defensible transcriptomic analysis must consequently separate directional concordance across bulk cohorts from differential expression within defined cell populations and from causal or therapeutic claims.",
        "Immune compartments create a parallel interpretive challenge. Alveolar and recruited macrophage populations differ in origin, state, lipid handling, and abundance in fibrotic lungs [9,26]. A macrophage-aggregate coefficient can therefore reflect a genuine within-state response, a change in the balance of macrophage subtypes, or both. Cell-type effects may be opposing rather than merely attenuated, making explicit subtype and absolute-expression analyses necessary before assigning a lung-wide direction.",
        "Cross-platform synthesis adds another layer of uncertainty. RNA-sequencing and microarray log fold changes are not directly exchangeable, and inference from only a few cohorts requires standardized effects, conservative random-effects confidence intervals, prediction intervals, and omission analyses [15]. Single-cell differential expression has a different replication unit: cells are nested within donors, making donor-level aggregation necessary to avoid pseudoreplication [16]. Technical libraries that are donor-specific and nested in disease cannot be treated as an ordinary batch covariate. Composition estimates are also reference-dependent and should be stress-tested rather than interpreted as measured cell fractions.",
        "We integrated source-level purified AT2, whole-lung, and lung single-cell transcriptomic analyses to ask whether a directionally concordant FASN signal represented a uniform intracellular decrease or a compartment-dependent tissue signal. An explicit reproducible AT2/pathway rule was applied to every qualifying gene rather than assumed to identify FASN uniquely. FASN was then used as the deepest contextual case because it had the strongest nominal whole-lung concordance among rule-positive genes and a direct role in de novo fatty-acid synthesis. The study was designed as a transparent contextual analysis, not an exhaustive metabolic target screen.",
    ], main_texts)

    doc.add_heading("Methods", level=1)
    doc.add_heading("Study design and public datasets", level=2)
    add_paragraphs(doc, [
        "This secondary study reanalyzed public transcriptomic data without generating new human or animal samples. The evidence sequence comprised purified AT2 pathway analysis, standardized whole-lung synthesis, full-transcriptome lung-cell pseudobulk, and composition-sensitive calibration. Only results regenerated from raw gene counts, raw CEL files, or the complete raw sparse UMI matrix were used as primary evidence. Dataset roles and boundaries are summarized in Table 1. GEO resources were obtained from the NCBI Gene Expression Omnibus [19].",
    ], main_texts)
    add_dataset_table(doc)

    doc.add_heading("Purified AT2 raw-count analysis", level=2)
    add_paragraphs(doc, [
        "GSE245965 was analyzed from the deposited AT2 gene-count table, comprising two IPF and three control samples [8]. Counts were filtered with edgeR, normalized by trimmed mean of M values, transformed with limma-voom, and modeled by disease status with robust empirical Bayes moderation [12-14]. Differential expression required BH-FDR<0.05 and |log2 fold change|>1 [18]. Quality control included library size, PCA, sample correlation, and the voom mean-variance trend.",
        "The moderated t statistic ranked genes for preranked enrichment. Hallmark and Reactome sets were obtained through MSigDB, and fgseaMultilevel was run with set sizes of 10-500 and BH correction [27,28]. Over-representation analysis used the same FDR-defined differential genes with clusterProfiler [11]. Leave-one-sample-out analyses repeated filtering, voom modeling, enrichment, and FASN extraction after omitting each sample. The focal-rule audit retained every gene that was FDR-significantly downregulated with |log2 fold change|>1, belonged to the leading edge of both Hallmark and Reactome fatty-acid metabolism, and was measurable in all three whole-lung cohorts. FASN was selected for deeper case analysis only after the complete qualifying set was reported.",
    ], main_texts)

    doc.add_heading("Whole-lung source processing and transcriptome-wide synthesis", level=2)
    add_paragraphs(doc, [
        "GSE150910 was reanalyzed from gene-level counts. Symbols were cleaned, duplicate genes summed, low-expression genes filtered, and TMM-voom expression modeled for 103 IPF and 103 control lungs [12-14]. GSE110147 was rebuilt from 48 raw CEL files using background correction, quantile normalization, and robust multi-array average at the core transcript level; 22 IPF and 11 controls were selected from GEO metadata. GSE24206 was rebuilt from 23 raw CEL files, yielding 17 IPF and 6 controls. Probe sets were mapped with GPL6244 or GPL570 annotations and collapsed by maximum interquartile range [29]. Deposited series matrices were retained only as preprocessing sensitivity analyses.",
        "For every gene measured in all three cohorts, the IPF-control contrast was standardized as Hedges' g with the small-sample correction and sampling variance. Restricted-maximum-likelihood random-effects models used Hartung-Knapp inference in metafor [15]. We report 95% confidence intervals, tau-squared, I2, prediction intervals, and BH-FDR across the complete three-cohort transcriptome intersection [18]. FASN was additionally evaluated by omitting each cohort in turn.",
    ], main_texts)

    doc.add_heading("Complete raw-matrix aggregation and donor pseudobulk", level=2)
    add_paragraphs(doc, [
        "GSE136831 comprised 45,947 genes, 312,928 cells, and 692,789,348 nonzero UMI coordinates with matched gene, barcode, donor, disease, cell-type, and Library_Identity metadata [9]. A streaming parser traversed every coordinate and aggregated counts to gene-by-donor-by-cell-type pseudobulk matrices. Control and IPF cells in 20 interpretable compartments contributed 232,056 cells and 964 donor-cell-type samples after alveolar and non-alveolar macrophages were separated. Raw-matrix column sums matched metadata total UMI values exactly in all 964 samples.",
        "As a source-integrity check, all coordinates for a precomputed 62-gene subset were reconstructed during the full-matrix scan. The original and reconstructed panels contained 2,363,040 entries and had the same canonical SHA-256 hash. This replaced the earlier limited first-5,000-cell comparison and confirmed the target subset across all 312,928 cells.",
        "Counts were analyzed separately by cell type with edgeR quasi-likelihood models, using donors as biological replicates and disease status as the predictor [12,16]. Genes were filtered with filterByExpr; FASN was forcibly retained as the explicitly selected focal case. Full-transcriptome total UMI supplied library sizes, and TMM factors were estimated after filtering while preserving those library sizes. The analysis required at least 5, 20, or 50 cells per donor-cell-type sample; 20 cells was the primary threshold and at least three donors per disease group were required. FASN P values were BH-corrected across estimable cell types within each threshold as the primary selected-gene family; full-transcriptome gene-level FDR within each cell type was secondary context. Figure confidence intervals are QL-based approximations reconstructed from the coefficient and moderated F statistic.",
        "Absolute-expression summaries included donor CP10K, detection rate, and normalized abundance among FASN-positive cells. Figure 4A reports the median and interquartile range across donors meeting the primary 20-cell threshold. Detection prevalence was modeled from donor-level detected and undetected cell counts with a quasibinomial generalized linear model; an unweighted empirical-logit model was retained as a secondary sensitivity. Positive-cell abundance was modeled as donor-level log2 mean cell-level CP10K among FASN-positive cells. P values were BH-corrected across estimable cell types within each model family. Alveolar macrophages (Manuscript_Identity Macrophage_Alveolar) were analyzed separately from non-alveolar macrophages (Manuscript_Identity Macrophage); finer state labels were not treated as independent biological populations because many were designated outliers or cell-cycle states.",
        "Library_Identity was audited against disease and donor. Each library belonged to one disease group and almost always one donor, while some donors contributed two to four libraries; no shared source-study or center variable was available. Consequently, a fixed-effect model containing both Library_Identity and disease was rank-deficient or biologically unidentified. We report disease-by-library tables and donor-pseudobulk MDS, and use a multiple-library donor indicator only as a technical-replication sensitivity. This indicator is not presented as source-library adjustment.",
    ], main_texts)

    doc.add_heading("Composition sensitivity", level=2)
    add_paragraphs(doc, [
        "A 44-45-marker reference matrix was generated from GSE136831 for major epithelial, stromal, immune, and endothelial compartments. Marker-score covariates and non-negative least-squares (NNLS) reference components were estimated after rescaling marker expression [17]. Whole-lung FASN expression was modeled unadjusted, with marker-score covariates, and with six NNLS components. These were treated as reference components rather than measured fractions.",
        "Stability analyses reported the reference-matrix condition number, variance-inflation factors (VIFs), 1,000 disease-stratified bootstrap confidence intervals, leave-one-marker-out coefficients, and leave-one-cell-type-out coefficients. Direction changes or high VIFs were interpreted as model dependence rather than evidence that composition caused the bulk difference.",
    ], main_texts)

    doc.add_heading("Statistical analysis and reproducibility", level=2)
    add_paragraphs(doc, [
        "Tests were two-sided and BH-FDR was applied to the complete family relevant to each analysis [18]. R 4.4.1 was used for differential expression, enrichment, RMA, meta-analysis, and composition models; Python and a streaming .NET parser were used for metadata preparation, sparse-matrix verification, and count aggregation [20]. Pipeline scripts accept raw-data and analysis roots as arguments; the document builders require only the analysis root. A manifest-driven downloader verifies the size and SHA-256 hash of every public input before analysis. Input hashes, software versions, source-data tables, and intermediate validation summaries accompany this version.",
    ], main_texts)

    doc.add_heading("Results", level=1)
    doc.add_heading("Study overview and reproducible focal selection", level=2)
    add_paragraphs(doc, [
        "The analysis linked four evidence levels: AT2 pathway discovery, transcriptome-wide whole-lung calibration, full-transcriptome donor pseudobulk, and reference-dependent composition sensitivity (Fig. 1). Six genes met the stated focal rule: FASN, ACADL, ACSL4, HSD17B4, ACSM3, and ACSL1. FASN was selected as the deepest contextual case because it had the smallest nominal whole-lung meta-analysis P among these six and a direct role in de novo fatty-acid synthesis. The whole-lung multiple-testing family was the 13,632 genes with estimable effects across all three cohorts.",
    ], main_texts)
    base.add_figure(doc, 1, "Study design and evidence progression. (A) A concordant whole-lung FASN decrease may represent intracellular regulation, altered composition, compartment-specific effects, or a combination. (B) Source-level datasets used in the primary analyses. (C) An explicit reproducible AT2/pathway rule yielded six genes; transcriptome-wide whole-lung calibration was then completed before FASN was selected for donor pseudobulk and composition-sensitive contextual analysis.")

    doc.add_heading("Purified AT2 cells coupled structural remodeling to suppressed lipid programs", level=2)
    add_paragraphs(doc, [
        "After filtering, 18,026 genes were retained. Robust limma-voom identified 789 upregulated and 795 downregulated genes at BH-FDR<0.05 and |log2 fold change|>1. FASN was downregulated (log2 fold change -3.512; P=0.00175; FDR=0.0332). PCA separated IPF from controls along PC1, although library sizes and correlations showed substantial sample-level variation (Supplementary Fig. S1).",
        "GSEA identified positive epithelial-mesenchymal transition (NES 2.315; FDR 1.08x10^-9) and extracellular-matrix organization (NES 2.153; FDR 3.22x10^-9), whereas cholesterol biosynthesis (NES -2.610; FDR 3.42x10^-7), fatty-acid metabolism (NES -2.172; FDR 6.20x10^-7), surfactant metabolism (NES -2.375; FDR 4.83x10^-5), unfolded-protein response (NES -1.941; FDR 9.05x10^-4), phospholipid metabolism, and oxidative phosphorylation were negatively enriched (Fig. 2). FASN belonged to the leading edge of both Hallmark and Reactome fatty-acid metabolism.",
        "Leave-one-sample-out analysis separated pathway stability from gene-list stability. Upregulated DEG counts ranged from 160 to 501 and downregulated counts from 126 to 520, but major structural, lipid, surfactant, and stress-response pathways retained their directions. Oxidative phosphorylation reversed direction when one IPF sample was omitted (Supplementary Fig. S2). FASN remained negative in all five omission models (log2 fold change -2.475 to -4.482), but its FDR ranged from 0.0347 to 0.127. We therefore interpret the pathway state and FASN direction more strongly than omission-specific significance from this five-sample cohort.",
    ], main_texts)
    base.add_figure(doc, 2, "Purified AT2 structural-metabolic divergence in two IPF and three control samples. (A) Differential expression from raw counts using BH-FDR<0.05 and |log2 fold change|>1. (B) Selected GSEA results; horizontal position is normalized enrichment score, color indicates direction, and point size represents -log10 FDR. AT2, alveolar epithelial type II; GSEA, gene-set enrichment analysis.")

    doc.add_heading("FASN was directionally concordant across whole-lung cohorts without transcriptome-wide discovery", level=2)
    add_paragraphs(doc, [
        "Source-level FASN effects were negative in GSE150910 (Hedges' g -0.750; 95% CI -1.032 to -0.467), GSE110147 (-1.561; 95% CI -2.382 to -0.739), and GSE24206 (-1.197; 95% CI -2.196 to -0.199). Raw-CEL RMA closely reproduced the deposited-series estimates in both microarray cohorts, indicating that the direction was not created by the deposited normalization (Fig. 3).",
        "The Hartung-Knapp pooled effect was -1.046 (95% CI -2.100 to 0.009; P=0.0508), with tau-squared 0.115, I2 49.0%, and a prediction interval of -2.847 to 0.756. Across 13,632 genes with estimable three-cohort effects, FASN had FDR 0.818 and no gene reached FDR<0.10. FASN ranked 825th by nominal P, within the top 6.1% but far from corrected discovery (Supplementary Fig. S3). Each leave-one-cohort-out FASN estimate was nonsignificant. Thus, concordant cohort directions did not constitute corrected transcriptome-wide discovery.",
    ], main_texts)
    base.add_figure(doc, 3, "Whole-lung FASN synthesis and transcriptome-wide calibration. (A) Cohort-specific FASN Hedges' g values and the REML Hartung-Knapp pooled estimate. (B) Transcriptome-wide standardized meta-analysis of 13,632 genes; none reached BH-FDR<0.10. (C) Raw-CEL RMA and deposited-series FASN effects. (D) Leave-one-cohort-out FASN estimates. CI, confidence interval; FDR, false-discovery rate.")

    doc.add_heading("Full-transcriptome pseudobulk localized opposing FASN effects across lung compartments", level=2)
    add_paragraphs(doc, [
        "FASN was most abundant in AT2 cells, with lower expression in AT1, airway epithelial, stromal, immune, and endothelial compartments (Fig. 4A). Figure 4A displays donor-level medians and interquartile ranges among donor-cell-type samples meeting the primary 20-cell threshold. At that threshold, nine cell types retained at least three control and three IPF donors for differential-expression modeling. AT1, basal, club, goblet, aberrant basaloid, and several rare populations did not meet primary-threshold donor requirements and were not assigned primary disease effects.",
        "Ciliated cells showed a negative FASN effect at the primary threshold (log2 fold change -1.234; approximate 95% CI -1.739 to -0.730; P=9.04x10^-6; FASN cell-type FDR 8.13x10^-5; gene-level FDR 0.00121; 13 control and 32 IPF donors). The direction persisted at the 5- and 50-cell thresholds (log2 fold changes -1.154 and -1.261). Median donor CP10K was 0.175 in controls and 0.0746 in IPF, with median detection rates of 6.82% and 5.43%. The donor-stratified quasibinomial model supported lower detection prevalence in IPF (odds ratio 0.489; P=4.94x10^-4; BH-FDR=0.00342), whereas normalized abundance among FASN-positive cells did not differ (P=0.532; FDR=0.598). The empirical-logit sensitivity had the same direction (P=0.00251; FDR=0.0150). Thus, the ciliated result is a threshold-consistent, low-abundance effect driven primarily by detection prevalence within GSE136831 (Supplementary Figs. S4-S5 and S7).",
        "Alveolar macrophages showed an opposing positive pseudobulk effect (log2 fold change 0.490; approximate 95% CI 0.194 to 0.785; P=0.00152; cell-type FDR 0.00686; gene-level FDR 0.00642; 26 control and 31 IPF donors). Median donor CP10K was 0.0270 in controls and 0.0314 in IPF, with median detection rates of 1.87% and 2.84%. Detection prevalence increased in the quasibinomial model (odds ratio 1.393; P=0.00296; FDR=0.00666), whereas positive-cell normalized abundance did not differ (P=0.342; FDR=0.513). Non-alveolar macrophages were not FDR-supported in pseudobulk (log2 fold change 0.200; P=0.135; cell-type FDR 0.160), despite opposing hurdle components of higher prevalence (FDR=0.00342) and lower positive-cell normalized abundance (FDR=0.00137). The aggregate macrophage observation was therefore refined to a low-abundance, compartment- and component-dependent pattern rather than a uniform macrophage increase.",
        "Monocytes were negative (log2 fold change -0.788; P=0.0126; cell-type FDR 0.0378), although their gene-level FDR was 0.0539; detection prevalence was also lower (odds ratio 0.524; P=0.0220; FDR=0.0331). AT2 cells showed no negative effect at the primary threshold (log2 fold change 0.602; P=0.142), and the AT1 estimate at the 5-cell threshold did not survive correction. Adding a multiple-library donor indicator retained the ciliated (-1.144; P=7.10x10^-5) and alveolar-macrophage (0.488; P=0.00154) directions. However, Library_Identity was donor-specific and fully nested within disease, and no shared source/center field was available; disease and source-library effects could not be separated. MDS displayed disease-associated structure, especially in macrophage compartments, but cannot resolve whether that structure is biological or source-related (Supplementary Fig. S8).",
    ], main_texts)

    doc.add_heading("Composition estimates were reference- and cohort-dependent", level=2)
    add_paragraphs(doc, [
        "The marker reference had a condition number near 5.1, but maximum VIF differed substantially: 3.28 in GSE150910, 7.81 in GSE24206, and 27.43 in GSE110147. Marker-score and NNLS adjustment changed the magnitude and sometimes the direction of FASN coefficients. All NNLS confidence intervals crossed zero, and GSE110147 estimates were particularly imprecise (Fig. 4C).",
        "Disease-stratified bootstrap intervals were narrowest in GSE150910 and wide in the two smaller cohorts. Leave-one-cell-type-out analysis changed coefficient direction in every cohort (Supplementary Fig. S6). Composition models therefore showed that the bulk estimate was sensitive to the chosen reference and covariates; they did not establish that altered composition quantitatively explained the FASN decrease.",
    ], main_texts)
    base.add_figure(doc, 4, "Cell-resolved and composition-sensitive interpretation of FASN. (A) Median donor FASN CP10K and interquartile range across cell compartments among samples meeting the primary 20-cell threshold; point size is the median donor percentage of cells with detected FASN. The CP10K axis uses a pseudo-log transformation to display low-abundance compartments. (B) Full-transcriptome edgeR FASN estimates at the primary threshold of at least 20 cells per donor-cell type; n denotes control/IPF donors, orange denotes BH-FDR<0.05 across estimable cell types, and labels identify the ciliated, alveolar-macrophage, and monocyte FDR values. (C) Whole-lung FASN coefficients before and after marker-score or NNLS adjustment. Adjusted values are reference-dependent sensitivity estimates, not measured cell fractions. NNLS, non-negative least squares.")

    doc.add_heading("Discussion", level=1)
    add_paragraphs(doc, [
        "This study shows that a directionally concordant whole-lung metabolic signal can remain statistically and biologically heterogeneous after source-level reanalysis. Purified AT2 cells exhibited a contrast between structural remodeling and suppressed lipid/surfactant programs, and FASN was reduced in all three independently processed whole-lung cohorts. Yet FASN did not reach transcriptome-wide meta-analysis FDR, its prediction interval crossed zero, and donor-level single-cell analysis identified a low-abundance negative ciliated-cell effect alongside a positive alveolar-macrophage prevalence signal. The most defensible conclusion is compartment-dependent remodeling within a single cell-resolved atlas, not a universal intracellular decrease or therapeutic target.",
        "Selection transparency materially changes the interpretation. The AT2 significance, dual leading-edge membership, and three-cohort measurability rule produced six genes, not FASN alone. FASN was prioritized because it had the smallest nominal whole-lung meta-analysis P among those six and a direct biochemical role in de novo fatty-acid synthesis. This is a defensible rationale for a contextual case study, but it is not evidence that FASN is the unique metabolic driver. ACADL, HSD17B4, ACSL1, ACSM3, and ACSL4 remain rule-positive observations, and the present work should not be read as an exhaustive target screen.",
        "The purified AT2 analysis supports coordinated metabolic stress but also illustrates the limits of a five-sample dataset. DEG counts varied several-fold after sample omission, whereas major extracellular-matrix, epithelial-transition, lipid, surfactant, and stress-response pathways largely retained direction. FASN remained negative after every omission, but its FDR crossed 0.05 in three of five omission models. This pattern favors a pathway-level epithelial state and directional sensitivity over a claim of sample-independent single-gene significance. Moreover, the AT2 decrease was not reproduced as a negative AT2 effect in GSE136831. Differences in purified-cell isolation, dissociation-based annotation, cell yield, cohort severity, and disease state may all contribute.",
        "The whole-lung result has a narrower meaning. Concordance across counts and independently normalized CEL files reduces concern about preprocessing artifacts, but it does not overcome uncertainty from only three cohorts. Hartung-Knapp inference widened the confidence interval, the prediction interval included positive effects, and no transcriptome-wide gene reached FDR<0.10. FASN can therefore be described as a concordant observation selected for contextual analysis, not as a replicated discovery.",
        "The airway epithelial finding is specific rather than generic. Ciliated cells showed a negative coefficient across all cell-count thresholds, while AT2 did not reproduce the purified-cell direction and several rarer epithelial populations were not estimable at the primary threshold. IPF lungs contain airway-like epithelial remodeling and bronchiolized regions alongside disrupted alveolar repair [9,23]. Reduced FASN detection in ciliated cells could accompany epithelial injury, altered differentiation, ciliary maintenance demands, or an airway-like remodeling state, but the current data cannot distinguish these possibilities. The result therefore localizes a statistical pattern; it does not establish a mechanism.",
        "Absolute abundance is central to that restraint. Median FASN detection among IPF ciliated-cell donors was 5.43%, with median CP10K below 0.1, and macrophage detection was lower still. Relative log-fold changes can be large when baseline counts are sparse and may be influenced by capture efficiency or ambient RNA. The donor-stratified quasibinomial analysis supported reduced ciliated detection prevalence, whereas normalized abundance among FASN-positive cells did not differ. This convergence suggests that fewer ciliated cells detect FASN in IPF, while leaving the magnitude of intracellular downregulation less secure.",
        "Macrophage separation also revised the story. The positive pseudobulk coefficient was supported in alveolar macrophages but not in the non-alveolar compartment, and the quasibinomial analysis implicated detection prevalence rather than higher normalized abundance among FASN-positive alveolar macrophages. The non-alveolar compartment showed opposing prevalence and positive-cell-abundance components, explaining why a single aggregate direction would be misleading. Fibrotic lungs contain resident, recruited, and disease-associated macrophage states with different lipid requirements and changing proportions [9,26]. Even the alveolar annotation may contain internal state mixtures. The counter-direction therefore illustrates how cell-state composition can generate opposing compartment signals, but it should not be generalized to every macrophage lineage.",
        "Technical structure remains the most important limitation of the single-cell layer. Library identifiers were unique or nearly unique to donors and were completely nested within disease, so a model containing both library and disease could not identify separate coefficients. The multiple-library indicator showed that technical replication density did not reverse the focal effects, but it cannot correct source-library confounding. Disease-associated MDS separation is consequently compatible with biological remodeling, technical structure, or both. The ciliated and alveolar-macrophage findings should be described as internally threshold-consistent within GSE136831, not externally replicated cell-type discoveries.",
        "The ciliated-cell result also should not be converted directly into a treatment claim. The analysis lacked spatial niche information, longitudinal sampling, perturbation, and an independent lung single-cell cohort. Spatial and cell-atlas studies demonstrate that fibrotic programs occupy structured epithelial and collagen-producing stromal niches that dissociation-based data cannot preserve [10,24]. A future validation should test whether ciliated FASN reduction and alveolar-macrophage prevalence changes recur in an independent atlas and colocalize with aberrant epithelial or fibrotic niches.",
        "Composition analyses reinforced rather than resolved this uncertainty. Adjusted coefficients depended on cohort, reference model, and included cell types; GSE110147 had severe collinearity, and leave-one-cell-type-out analyses changed direction. The appropriate interpretation is model sensitivity. Independent single-cell references and methods such as MuSiC or BisqueRNA would help determine whether any adjusted direction is robust, but even a stable deconvolution result would remain an estimate rather than direct measurement.",
        "These distinctions matter because lipid synthesis can support epithelial membrane, surfactant, and stress-response biology while also sustaining activated stromal programs [5-7,22]. A lung-wide intervention inferred from a bulk direction could therefore have different consequences across epithelial, stromal, and immune compartments. The current data do not identify whether increasing or inhibiting FASN would be beneficial in vivo. They instead define an evidence sequence for future metabolic studies: source-level reprocessing, transcriptome-wide calibration, donor-level subtype analysis, composition stress tests, and spatial or perturbational validation.",
        "Several limitations remain. GSE245965 includes only five samples, and FASN omission-model FDR was sensitive to individual samples. The whole-lung meta-analysis contains three cohorts and estimates heterogeneity imprecisely. GSE136831 is a single dissociation-based dataset in which source-library and disease are inseparable; age, sex, smoking, treatment, and shared center variables were unavailable for fitted adjustment. Rare epithelial subtypes lost control-donor support, low FASN abundance increases measurement uncertainty, and edgeR confidence intervals are QL-based approximations. The composition analysis used one reference and approximate NNLS or marker-score components rather than measured fractions. No independent cell-resolved, spatial, or wet-laboratory validation was available. Within these limits, source-level analysis supports a cell-, prevalence-, and reference-dependent interpretation of the concordant whole-lung FASN signal.",
    ], main_texts)

    doc.add_heading("Declarations", level=1)
    for label, text in [
        ("Author contributions", "A final CRediT statement must be confirmed by all authors before submission."),
        ("Funding", "Funding information must be completed before submission."),
        ("Competing interests", "Author-specific declarations must be completed before submission."),
        ("Ethics approval", "This secondary analysis used de-identified public datasets. Ethics approval and informed consent were reported by the original studies; no new participants were enrolled."),
        ("Consent for publication", "Not applicable to this secondary analysis of de-identified public data."),
        ("Data availability", "Public data are available under GEO accessions GSE245965, GSE150910, GSE110147, GSE24206, and GSE136831. Input hashes and source-data tables accompany this version. A permanent public repository DOI should be added before submission."),
        ("Code availability", "Versioned R, Python, and streaming aggregation code accompany this version. A public archive should be created before submission."),
        ("Acknowledgements", "An acknowledgement statement must be confirmed before submission."),
        ("Supplementary material", "The supplement contains eight supplementary figures and a workbook with source verification, focal-rule transparency, AT2 sensitivity, transcriptome-wide whole-lung results, full-transcriptome pseudobulk, library audit, low-abundance sensitivity, and composition diagnostics."),
    ]:
        base.add_labeled_paragraph(doc, label + ":", text)

    doc.add_heading("References", level=1)
    for reference in references():
        paragraph = doc.add_paragraph(style="Reference")
        paragraph.paragraph_format.line_spacing = 0.92
        paragraph.paragraph_format.space_after = Pt(0)
        base.set_font(paragraph.add_run(reference), size=8.5)

    abstract_count = count_words([f"{label}: {text}" for label, text in abstract_data])
    main_count = count_words(main_texts)
    for paragraph in doc.paragraphs:
        if "WORD_COUNTS_PLACEHOLDER" in paragraph.text:
            paragraph.clear()
            base.set_font(paragraph.add_run(
                f"Word counts: Abstract {abstract_count}; main text {main_count} | Tables 1 | Figures 4 | Supplementary figures 8"
            ), size=9.5, color=base.MUTED)

    doc.core_properties.title = "Cell-Resolved Reanalysis Reframes a Concordant Whole-Lung FASN Signal in Idiopathic Pulmonary Fibrosis"
    doc.core_properties.subject = "Source-level cross-cohort and full-transcriptome pseudobulk reanalysis"
    doc.core_properties.keywords = "IPF; FASN; ciliated epithelium; pseudobulk; tissue composition"
    doc.core_properties.author = "Sutong Chen et al."
    doc.save(OUTPUT)
    return abstract_count, main_count


def build_action_items() -> None:
    doc = Document()
    base.configure_document(doc)
    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.text = "Cell-resolved reanalysis of FASN in IPF"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in header.runs:
            base.set_font(run, size=8.5, italic=True, color=base.MUTED)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    base.set_font(title.add_run("Remaining Submission Actions After v1.3.1 Release Finalization"), size=16, bold=True, color=RGBColor(22, 41, 55))
    lead = doc.add_paragraph()
    lead.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    base.set_font(lead.add_run("Focal-rule reconstruction, AT2 FASN omission analysis, source-library audit, macrophage subtype separation, low-abundance sensitivity, figures, and manuscript narrative are complete. Remaining P0 items require author or repository-owner input."), size=10.5)
    items = [
        ("P0", "Confirm author statements", "Complete CRediT contributions, funding, competing interests, acknowledgements, ORCID iDs, degrees, and official English affiliations.", "Open"),
        ("P0", "Create a permanent archive", "Deposit code, hashes, source data, software versions, workbook, and figures in a public repository and mint a DOI.", "Open"),
        ("P0", "Lock the master version", "Use Manuscript_v1.3.1_20260712, Supplement_v1.3.1_20260712, and Figures_v1.3.1_20260712 only.", "Open"),
        ("Future", "Independent cell-context validation", "An independent atlas, spatial analysis, or cell-specific perturbation could raise the evidence ceiling but is not required for the present v1.3.1 package.", "Deferred"),
    ]
    table = doc.add_table(rows=1, cols=4)
    for index, value in enumerate(["Priority", "Item", "Required action", "Status"]):
        table.rows[0].cells[index].text = value
        base.set_cell_shading(table.rows[0].cells[index], base.HEADER_FILL)
        for paragraph in table.rows[0].cells[index].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                base.set_font(run, size=8.8, bold=True)
    base.set_repeat_table_header(table.rows[0])
    for row_index, (priority, item, action, status) in enumerate(items, start=1):
        cells = table.add_row().cells
        for column, value in enumerate([priority, item, action, status]):
            cells[column].text = value
            cells[column].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if row_index % 2 == 0:
                base.set_cell_shading(cells[column], "F7F9FB")
            for paragraph in cells[column].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    base.set_font(run, size=8.3, bold=(column == 0))
    base.set_table_geometry(table, [0.55, 1.35, 3.95, 0.55])
    doc.save(ACTION_OUTPUT)


def main() -> None:
    for path in FIGURES.values():
        if not path.exists():
            raise FileNotFoundError(path)
    abstract_count, main_count = build_manuscript()
    build_action_items()
    print(f"Wrote {OUTPUT}")
    print(f"Wrote {ACTION_OUTPUT}")
    print(f"Abstract words: {abstract_count}; main-text words: {main_count}")


if __name__ == "__main__":
    main()
